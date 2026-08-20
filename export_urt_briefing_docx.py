# -*- coding: utf-8 -*-
"""将 URT 领导汇报版 Markdown 导出为严肃汇报体例 Word。

排版约定（业界常见中文汇报稿）：
- 封面标题：黑体 22pt 加粗，居中
- 一级章节（一、二…）：黑体 16pt 加粗
- 二级小节（x.x）：黑体 14pt 加粗
- 正文：宋体 12pt，1.5 倍行距，段后 6pt
- 表格：表头加粗、浅灰底；单元格宋体 10.5pt，垂直居中
- 页边距：上下 2.54cm，左右 3.17cm
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "docs" / "strategies" / "urt" / "URT上升趋势策略说明书_领导汇报版.md"
OUT_PATH = ROOT / "exported_docs" / "URT上升趋势策略说明书_领导汇报版.docx"

FONT_TITLE = "黑体"
FONT_HEADING = "黑体"
FONT_BODY = "宋体"
FONT_CODE = "Consolas"


def set_run_font(run, *, name_cn: str, size_pt: float, bold: bool = False, color: RGBColor | None = None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = name_cn
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name_cn if name_cn != FONT_BODY else "Times New Roman")
    rFonts.set(qn("w:hAnsi"), name_cn if name_cn != FONT_BODY else "Times New Roman")
    rFonts.set(qn("w:eastAsia"), name_cn)
    rFonts.set(qn("w:cs"), name_cn if name_cn != FONT_BODY else "Times New Roman")


def set_paragraph_format(p, *, line_spacing: float = 1.5, space_before: float = 0, space_after: float = 6, first_line: float | None = None):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)


def add_styled_paragraph(
    doc: Document,
    text: str,
    *,
    style: str = "body",
    align: WD_ALIGN_PARAGRAPH | None = None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align

    if style == "title":
        set_paragraph_format(p, line_spacing=1.75, space_before=0, space_after=20)
        run = p.add_run(text)
        set_run_font(run, name_cn=FONT_TITLE, size_pt=22, bold=True)
    elif style == "h1":
        set_paragraph_format(p, line_spacing=1.75, space_before=20, space_after=12)
        run = p.add_run(text)
        set_run_font(run, name_cn=FONT_HEADING, size_pt=16, bold=True)
    elif style == "h2":
        set_paragraph_format(p, line_spacing=1.75, space_before=14, space_after=10)
        run = p.add_run(text)
        set_run_font(run, name_cn=FONT_HEADING, size_pt=14, bold=True)
    elif style == "h3":
        set_paragraph_format(p, line_spacing=1.75, space_before=10, space_after=8)
        run = p.add_run(text)
        set_run_font(run, name_cn=FONT_HEADING, size_pt=12, bold=True)
    elif style == "note":
        set_paragraph_format(p, line_spacing=1.5, space_before=8, space_after=8)
        run = p.add_run(text)
        set_run_font(run, name_cn=FONT_BODY, size_pt=10.5, bold=False, color=RGBColor(0x55, 0x55, 0x55))
    else:
        set_paragraph_format(p, line_spacing=1.75, space_before=0, space_after=10, first_line=0.74)
        _add_inline_runs(p, text, size_pt=12)

    return p


def _add_inline_runs(p, text: str, *, size_pt: float = 12, bullet: bool = False):
    """支持 **加粗** 片段。"""
    if bullet:
        set_paragraph_format(p, line_spacing=1.75, space_before=3, space_after=6)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, name_cn=FONT_BODY, size_pt=size_pt, bold=True)
        else:
            # 清理行内代码反引号，保留内容
            clean = part.replace("`", "")
            run = p.add_run(clean)
            set_run_font(run, name_cn=FONT_BODY, size_pt=size_pt, bold=False)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    # 移除样式自带 run，统一用正文宋体重写
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    _add_inline_runs(p, text, size_pt=12, bullet=True)
    return p


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing=1.35, space_before=0, space_after=2)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line.strip() else " ")
        set_run_font(run, name_cn=FONT_CODE, size_pt=10, bold=False)
        # eastAsia 仍给宋体，避免中文乱码
        rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), FONT_BODY)
        rFonts.set(qn("w:ascii"), FONT_CODE)
        rFonts.set(qn("w:hAnsi"), FONT_CODE)


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")


def set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 10.5, header: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, line_spacing=1.35, space_before=3, space_after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    # 支持单元格内 **加粗**
    parts = re.split(r"(\*\*[^*]+\*\*)", text.replace("<br>", "\n").replace("&lt;", "<").replace("&gt;", ">"))
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, name_cn=FONT_BODY, size_pt=size_pt, bold=True)
        else:
            run = p.add_run(part.replace("`", ""))
            set_run_font(run, name_cn=FONT_BODY, size_pt=size_pt, bold=bold or header)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(cols):
            val = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            set_cell_text(cell, val, header=(i == 0), bold=(i == 0), size_pt=10.5)
            if i == 0:
                shade_cell(cell, "D9E2F3")
    # 表后空行
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.0, space_before=0, space_after=6)


def parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        # skip separator
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def configure_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def build():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_page(doc)

    # 默认正文样式
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(12)
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    except Exception:
        pass

    i = 0
    title_done = False
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
                i += 1
                continue
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("```"):
            in_code = True
            code_lines = []
            i += 1
            continue

        if line.startswith("# ") and not title_done:
            add_styled_paragraph(doc, line[2:].strip(), style="title", align=WD_ALIGN_PARAGRAPH.CENTER)
            title_done = True
            i += 1
            continue

        if line.startswith("## "):
            add_styled_paragraph(doc, line[3:].strip(), style="h1")
            i += 1
            continue

        if line.startswith("### "):
            add_styled_paragraph(doc, line[4:].strip(), style="h2")
            i += 1
            continue

        if line.startswith("#### "):
            add_styled_paragraph(doc, line[5:].strip(), style="h3")
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, next_i = parse_table_block(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        if line.lstrip().startswith("- "):
            add_bullet(doc, line.lstrip()[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line.lstrip()):
            # 有序列表按正文处理，保留序号
            p = doc.add_paragraph()
            set_paragraph_format(p, line_spacing=1.5, space_before=2, space_after=4)
            _add_inline_runs(p, line.strip(), size_pt=12)
            i += 1
            continue

        # 编制说明等灰色注记
        if line.startswith("**编制说明**") or line.startswith("注："):
            add_styled_paragraph(doc, line.replace("**", ""), style="note")
            i += 1
            continue

        add_styled_paragraph(doc, line.strip(), style="body")
        i += 1

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT_PATH)
        print(str(OUT_PATH))
    except PermissionError:
        alt = OUT_PATH.with_name(f"URT上升趋势策略说明书_领导汇报版_{Path(__file__).stem}.docx")
        # fallback timestamp
        from datetime import datetime

        alt = OUT_PATH.with_name(
            f"URT上升趋势策略说明书_领导汇报版_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        doc.save(alt)
        print(f"原文件被占用，已另存为: {alt}")


if __name__ == "__main__":
    build()
